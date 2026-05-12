"""
Парсер учебного плана (УП.xlsx) + рабочих программ дисциплин (РПД.docx).

Структура РПД (один docx — все дисциплины):
  Дисциплина
    └─ Компетенция      (ур. 1, напр. «УК-7 - Способен…»)
         └─ Индикатор   (ур. 2, напр. «УК-7.1 : Планирует…»)
              └─ Знать / Уметь / Владеть   (ур. 3)
  + итоговый блок «В РЕЗУЛЬТАТЕ ОСВОЕНИЯ» с суммарными ЗУВ

Использование:
  python parse_up_rpd.py --up УП.xlsx --rpd РПД.docx --out output/

Зависимости:
  pip install openpyxl python-docx
"""

import json, re, argparse, difflib
from pathlib import Path


# ─────────────────────────────────────────────────────────────────────────────
# Вспомогательные функции
# ─────────────────────────────────────────────────────────────────────────────

def _clean(t):
    return re.sub(r"[ \t]+", " ", t.replace("\n", " ").replace("\r", "")).strip()

def _cell_texts_fast(table) -> list:
    seen, out = set(), []
    for row in table.rows:
        for cell in row.cells:
            t = _clean(cell.text)
            if t and t not in seen:
                seen.add(t); out.append(t)
    deduped = []
    prev = None
    for t in out:
        if t != prev:
            deduped.append(t)
        prev = t
    return deduped

def _cell_texts_full(table) -> list:
    out = []
    for row in table.rows:
        seen_row = set()
        for cell in row.cells:
            t = _clean(cell.text)
            if t and t not in seen_row:
                seen_row.add(t); out.append(t)
    deduped = []
    prev = None
    for t in out:
        if t != prev:
            deduped.append(t)
        prev = t
    return deduped

def _znu_from_lines(lines: list) -> dict:
    """Извлекает Знать / Уметь / Владеть из списка строк."""
    result = {"Знать": [], "Уметь": [], "Владеть": []}
    cur = None
    for line in lines:
        low = line.lower().lstrip("-– \t")
        if low.startswith("знать"):
            cur = "Знать"
            tail = re.sub(r"^знать\s*:?\s*", "", line, flags=re.I).strip("-– ")
            if tail: result[cur].append(tail)
        elif low.startswith("уметь"):
            cur = "Уметь"
            tail = re.sub(r"^уметь\s*:?\s*", "", line, flags=re.I).strip("-– ")
            if tail: result[cur].append(tail)
        elif low.startswith("владеть"):
            cur = "Владеть"
            tail = re.sub(r"^владеть\s*:?\s*", "", line, flags=re.I).strip("-– ")
            if tail: result[cur].append(tail)
        elif cur:
            result[cur].append(line.strip("-– "))
    return {k: " ".join(v).strip() for k, v in result.items() if v}

def _extract_code_desc(line):
    """Из «УК-7 - Описание» возвращает («УК-7», «Описание»)."""
    m = re.match(r"^([А-ЯёA-Z\-]+\d+(?:\.\d+)?)\s*[-–:]\s*(.*)", line, re.I)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return line[:20].strip(), line.strip()

def _base_code(code):
    """«УК-7.1» → «УК-7» (для дедупликации блоков компетенций)."""
    return re.sub(r"\.\d+$", "", code).upper()


# ─────────────────────────────────────────────────────────────────────────────
# Регулярки
# ─────────────────────────────────────────────────────────────────────────────

_RE_IND    = re.compile(r"^[А-ЯёA-Z\-]+\d+\.\d+\s*[:–-]", re.I)
_RE_COMP   = re.compile(r"^[А-ЯёA-Z\-]+\d+\s*[:–-]", re.I)
_RE_ZNU    = re.compile(r"^(знать|уметь|владеть)\s*:?\s*", re.I)
_RE_TITLE  = re.compile(r"Рабочая программа дисциплины", re.I)
_RE_TITLE2 = re.compile(r"Рабочая программа", re.I)
_RE_S3     = re.compile(r"КОМПЕТЕНЦИИ ОБУЧАЮЩЕГОСЯ", re.I)
_RE_S3B    = re.compile(r"ПЛАНИРУЕМЫЕ РЕЗУЛЬТАТЫ", re.I)
_RE_RESULT = re.compile(r"В РЕЗУЛЬТАТЕ ОСВОЕНИЯ ДИСЦИПЛИНЫ.*ДОЛЖЕН", re.I)
_RE_STRUCT = re.compile(r"СТРУКТУРА И СОДЕРЖАНИЕ", re.I)
_RE_SKIP   = re.compile(
    r"(УП:\s*\d|стр\.\s*\d|"
    r"ХАРАКТЕРИЗУЮЩИЕ|овладеть компетенциями)", re.I
)


# ─────────────────────────────────────────────────────────────────────────────
# Класс-парсер одной дисциплины
# ─────────────────────────────────────────────────────────────────────────────

class _DiscParser:
    """
    Конечный автомат для извлечения компетенций/индикаторов/ЗУВ.

    Ключевое решение по дубликатам:
      В РПД каждая компетенция встречается дважды:
        1) в списке «УК-7 - Способен…»           ← первое вхождение
        2) в блоке «ПЛАНИРУЕМЫЕ РЕЗУЛЬТАТЫ»       ← второе (здесь индикаторы!)
      При втором вхождении мы НЕ создаём новый объект компетенции,
      а продолжаем дописывать индикаторы в уже существующий.
    """

    def __init__(self):
        self._competencies = []
        self._comp_by_base = {}  # base_code -> comp dict
        self._cur_comp = None
        self._cur_ind = None
        self._znu_buf = []
        self._znu_target = None
        self._sum_lines = []

    # ── helpers ─────────────────────────────────────────────────────────────

    def _flush_znu(self):
        if not self._znu_buf: return
        outcomes = _znu_from_lines(self._znu_buf)
        if self._znu_target == "ind" and self._cur_ind is not None:
            self._cur_ind["outcomes"].update(outcomes)
        elif self._znu_target == "sum":
            self._sum_lines += self._znu_buf
        self._znu_buf.clear()
        self._znu_target = None

    def _save_ind(self):
        if self._cur_ind and self._cur_comp:
            existing = {i["code"] for i in self._cur_comp["indicators"]}
            if self._cur_ind["code"] not in existing:
                self._cur_comp["indicators"].append(self._cur_ind)
        self._cur_ind = None

    def _save_comp(self):
        if self._cur_comp and self._cur_comp not in self._competencies:
            self._competencies.append(self._cur_comp)
        self._cur_comp = None

    # ── public ──────────────────────────────────────────────────────────────

    def feed(self, line: str):
        # Итоговый блок «В РЕЗУЛЬТАТЕ ОСВОЕНИЯ»
        if _RE_RESULT.search(line):
            self._flush_znu(); self._save_ind()
            self._znu_target = "sum"; self._znu_buf = []; return

        # Технические строки — пропускаем
        if _RE_SKIP.search(line): return

        # Индикатор (УК-7.1, ОПК-3.2, …) — обрабатываем ВСЕГДА
        if _RE_IND.match(line):
            self._flush_znu(); self._save_ind()
            code, desc = _extract_code_desc(line)
            self._cur_ind = {"code": code, "description": desc, "outcomes": {}}
            # Убеждаемся, что cur_comp указывает на правильную компетенцию
            base = _base_code(code)
            if base in self._comp_by_base:
                self._cur_comp = self._comp_by_base[base]
            self._znu_target = "ind"; self._znu_buf = []; return

        # Компетенция (УК-7, ОПК-3, …)
        if _RE_COMP.match(line):
            self._flush_znu(); self._save_ind(); self._save_comp()
            code, desc = _extract_code_desc(line)
            base = _base_code(code)
            if base in self._comp_by_base:
                # Дубликат — переключаем cur_comp на существующий объект
                self._cur_comp = self._comp_by_base[base]
            else:
                # Новая компетенция
                comp = {"code": code, "description": desc, "indicators": []}
                self._comp_by_base[base] = comp
                self._cur_comp = comp
            self._znu_target = None; self._znu_buf = []; return

        # Знать / Уметь / Владеть
        if _RE_ZNU.match(line) or (self._znu_target and line.startswith("-")):
            self._znu_buf.append(line); return

        if self._znu_target:
            self._znu_buf.append(line)

    def finalize(self):
        self._flush_znu(); self._save_ind(); self._save_comp()
        # Итоговые ЗУВ
        summary = _znu_from_lines(self._sum_lines) if self._sum_lines else {}
        # Финальный список компетенций (в порядке появления)
        comps = list(self._comp_by_base.values())
        return comps, summary


# ─────────────────────────────────────────────────────────────────────────────
# 1. ПАРСИНГ РПД
# ─────────────────────────────────────────────────────────────────────────────

def parse_rpd(docx_path, verbose=False):
    try:
        from docx import Document
    except ImportError:
        raise ImportError("pip install python-docx")

    doc = Document(docx_path)
    tables = doc.tables
    n = len(tables)

    disc_starts = []
    for i, t in enumerate(tables):
        texts = _cell_texts_fast(t)
        found_title = False
        for tx in texts:
            if _RE_TITLE.search(tx) or _RE_TITLE2.search(tx):
                found_title = True; continue
            if found_title and tx and not re.match(
                    r"(УП:|стр\.|УТВЕРЖДАЮ|И\.о\.|Минобр|кафедра|Рабочая программа)", tx, re.I):
                disc_starts.append((i, tx)); break

    if not disc_starts:
        if verbose:
            print("    [debug] Стандартный поиск не дал результатов, пробуем по абзацам...")
        for i, p in enumerate(doc.paragraphs):
            t = _clean(p.text)
            if not t:
                continue
            if _RE_TITLE.search(t) or _RE_TITLE2.search(t):
                for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                    nt = _clean(doc.paragraphs[j].text)
                    if nt and not re.match(r"(УП:|стр\.|УТВЕРЖДАЮ|кафедра|Рабочая программа)", nt, re.I):
                        disc_starts.append((-1, nt))
                        break

    seen_n, unique = set(), []
    for idx, name in disc_starts:
        if name not in seen_n:
            seen_n.add(name); unique.append((idx, name))

    if verbose:
        print(f"    Дисциплин в РПД: {len(unique)}")
        for _, nm in unique:
            print(f"      — {nm}")

    rpd_map = {}

    for pos, (start_i, disc_name) in enumerate(unique):
        if start_i < 0:
            continue
        end_i = unique[pos + 1][0] if pos + 1 < len(unique) else n

        in_s3, lines = False, []
        for t in tables[start_i:end_i]:
            for tx in _cell_texts_full(t):
                if _RE_S3.search(tx) or _RE_S3B.search(tx):
                    in_s3 = True; continue
                if _RE_STRUCT.search(tx):
                    in_s3 = False; continue
                if in_s3:
                    lines.append(tx)

        if not lines and verbose:
            print(f"    [debug] Нет строк для «{disc_name}»")

        parser = _DiscParser()
        for line in lines:
            parser.feed(line)
        comps, summary = parser.finalize()

        rpd_map[disc_name] = {"competencies": comps, "summary_outcomes": summary}

    if not rpd_map and verbose:
        print("    [debug] РПД-мап пуст, пробуем парсинг всех таблиц целиком...")
        all_lines = []
        in_s3 = False
        for t in tables:
            for tx in _cell_texts_full(t):
                if _RE_S3.search(tx) or _RE_S3B.search(tx):
                    in_s3 = True; continue
                if _RE_STRUCT.search(tx):
                    in_s3 = False; continue
                if in_s3:
                    all_lines.append(tx)
        if all_lines:
            parser = _DiscParser()
            for line in all_lines:
                parser.feed(line)
            comps, summary = parser.finalize()
            rpd_map["_all_"] = {"competencies": comps, "summary_outcomes": summary}

    return rpd_map


# ─────────────────────────────────────────────────────────────────────────────
# 2. ПАРСИНГ УЧЕБНОГО ПЛАНА
# ─────────────────────────────────────────────────────────────────────────────

def parse_up(xlsx_path):
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError("pip install openpyxl")

    wb = load_workbook(xlsx_path, read_only=True, data_only=True)
    if "ПланСвод" not in wb.sheetnames:
        raise ValueError(f"Лист 'ПланСвод' не найден. Листы: {wb.sheetnames}")
    ws = wb["ПланСвод"]

    SKIP = ["Блок", "Обязательная часть", "Часть, формируемая",
            "ФТД", "Практика", "Государственная", "Считать"]
    result = []
    for row in ws.iter_rows(min_row=4, values_only=True):
        marker, name, hours = row[0], row[1], row[10]
        if marker not in ("+", "-") or not isinstance(name, str): continue
        name = name.strip()
        if any(k in name for k in SKIP): continue
        try:    h = int(hours) if hours else 0
        except: h = 0
        result.append({"name": name, "hours": h, "active": marker == "+"})

    wb.close()
    return result


# ─────────────────────────────────────────────────────────────────────────────
# 3. ОБЪЕДИНЕНИЕ
# ─────────────────────────────────────────────────────────────────────────────

def _norm(s):
    return re.sub(r"\s+", " ", s.lower().strip())

def _fuzzy_match(name, candidates, threshold=0.82):
    nl = _norm(name)
    for key in candidates:
        kl = _norm(key)
        if kl == nl:
            return key
        if kl in nl or nl in kl:
            return key
    best_score, best_key = 0, None
    for key in candidates:
        kl = _norm(key)
        score = difflib.SequenceMatcher(None, nl, kl).ratio()
        if score > best_score:
            best_score, best_key = score, key
    if best_score >= threshold:
        return best_key
    return None

def build_combined(disciplines, rpd_map, verbose=False):
    idx = {_norm(k): v for k, v in rpd_map.items()}
    rpd_keys = list(rpd_map.keys())
    result = []
    for d in disciplines:
        dn = _norm(d["name"])
        rpd_data = idx.get(dn)
        if not rpd_data:
            fk = _fuzzy_match(d["name"], rpd_keys)
            if fk:
                rpd_data = rpd_map[fk]
                if verbose:
                    print(f"    [fuzzy] «{d['name']}» ↔ «{fk}»")
        result.append({
            "name":             d["name"],
            "hours":            d["hours"],
            "active":           d["active"],
            "competencies":     (rpd_data or {}).get("competencies", []),
            "summary_outcomes": (rpd_data or {}).get("summary_outcomes", {}),
        })
    return result


# ─────────────────────────────────────────────────────────────────────────────
# 4. ЭКСПОРТ
# ─────────────────────────────────────────────────────────────────────────────

def save_json(data, path):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"[OK] JSON: {path}")


def save_report(data, path):
    lines = [
        "=" * 72, "ОТЧЁТ ПО ДИСЦИПЛИНАМ УЧЕБНОГО ПЛАНА", "=" * 72,
        f"Всего: {len(data)}, активных: {sum(1 for d in data if d['active'])}", "",
    ]
    for i, d in enumerate(data, 1):
        mark = "+" if d["active"] else "-"
        lines.append(f"{i:>3}. [{mark}] {d['name']}  ({d['hours']} ч.)")
        comps = d.get("competencies", [])
        if comps:
            for comp in comps:
                lines.append(f"       ├─ [{comp['code']}] {comp['description']}")
                for ind in comp.get("indicators", []):
                    lines.append(f"       │   ├─ [{ind['code']}] {ind['description']}")
                    for k, v in ind.get("outcomes", {}).items():
                        lines.append(f"       │   │     {k}: {v}")
        else:
            lines.append("       └─ (РПД не найдена)")
        so = d.get("summary_outcomes", {})
        if so:
            lines.append("       └─ ИТОГО:")
            for k, v in so.items():
                lines.append(f"              {k}: {v}")
        lines.append("")

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"[OK] Отчёт: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# 5. CLI
# ─────────────────────────────────────────────────────────────────────────────

def main():
    p = argparse.ArgumentParser(description="Парсер УП.xlsx + РПД.docx")
    p.add_argument("--up",  required=True, help="Путь к УП.xlsx")
    p.add_argument("--rpd", help="Путь к РПД.docx")
    p.add_argument("--out", default=".", help="Папка для результатов")
    p.add_argument("--verbose", action="store_true", help="Подробный вывод")
    args = p.parse_args()

    out = Path(args.out)
    out.mkdir(parents=True, exist_ok=True)

    print(f"[→] УП: {args.up}")
    disciplines = parse_up(args.up)
    print(f"    Дисциплин: {len(disciplines)}")

    rpd_map = {}
    if args.rpd:
        print(f"[→] РПД: {args.rpd}")
        rpd_map = parse_rpd(args.rpd, verbose=args.verbose)
        matched = sum(1 for d in disciplines if _norm(d["name"]) in {_norm(k) for k in rpd_map})
        if matched == 0 and rpd_map:
            fk_count = 0
            for d in disciplines:
                if _fuzzy_match(d["name"], rpd_map):
                    fk_count += 1
            print(f"    Точных совпадений: 0, fuzzy-совпадений: {fk_count}")
        else:
            print(f"    Сопоставлено: {matched}/{len(disciplines)}")

    combined = build_combined(disciplines, rpd_map, verbose=args.verbose)
    save_json(combined, str(out / "disciplines_data.json"))
    save_report(combined, str(out / "disciplines_report.txt"))


if __name__ == "__main__":
    main()
